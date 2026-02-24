"""Document processors for various file types."""

import mimetypes
from dataclasses import dataclass
from pathlib import Path

import chardet

# Optional imports - gracefully handle missing dependencies
try:
    from docx import Document as DocxDocument

    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from pypdf import PdfReader

    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    import pytesseract
    from PIL import Image

    HAS_OCR = True
except ImportError:
    HAS_OCR = False


@dataclass
class ProcessingResult:
    """Result of processing a document."""

    text: str
    file_type: str
    success: bool
    error: str | None = None


# File extensions mapped to their types
TEXT_EXTENSIONS = {
    ".txt",
    ".md",
    ".markdown",
    ".rst",
    ".json",
    ".csv",
    ".xml",
    ".yaml",
    ".yml",
    ".toml",
    ".ini",
    ".cfg",
    ".conf",
    ".log",
    ".html",
    ".htm",
    ".css",
}

CODE_EXTENSIONS = {
    ".py",
    ".js",
    ".ts",
    ".jsx",
    ".tsx",
    ".java",
    ".c",
    ".cpp",
    ".cc",
    ".cxx",
    ".h",
    ".hpp",
    ".cs",
    ".go",
    ".rs",
    ".rb",
    ".php",
    ".swift",
    ".kt",
    ".kts",
    ".scala",
    ".r",
    ".R",
    ".sql",
    ".sh",
    ".bash",
    ".zsh",
    ".ps1",
    ".bat",
    ".cmd",
    ".lua",
    ".pl",
    ".pm",
    ".tcl",
    ".vim",
    ".el",
    ".clj",
    ".cljs",
    ".ex",
    ".exs",
    ".erl",
    ".hrl",
    ".hs",
    ".ml",
    ".mli",
    ".fs",
    ".fsx",
    ".dart",
    ".vue",
    ".svelte",
}

IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".tiff", ".tif", ".bmp", ".gif", ".webp"}


class DocumentProcessor:
    """Process various document types and extract text."""

    def __init__(self):
        """Initialize the processor."""
        # Initialize mimetypes
        mimetypes.init()

    def get_file_type(self, file_path: Path) -> str:
        """Determine the file type category."""
        ext = file_path.suffix.lower()

        if ext in TEXT_EXTENSIONS:
            return "text"
        elif ext in CODE_EXTENSIONS:
            return "code"
        elif ext == ".pdf":
            return "pdf"
        elif ext == ".docx":
            return "docx"
        elif ext in IMAGE_EXTENSIONS:
            return "image"
        elif ext == ".zip":
            return "archive"
        else:
            # Try to guess if it's text based on mimetype
            mime_type, _ = mimetypes.guess_type(str(file_path))
            if mime_type and mime_type.startswith("text/"):
                return "text"
            return "unknown"

    def process(self, file_path: Path) -> ProcessingResult:
        """Process a file and extract text.

        Args:
            file_path: Path to the file

        Returns:
            ProcessingResult with extracted text
        """
        file_type = self.get_file_type(file_path)

        try:
            if file_type in ("text", "code"):
                return self._process_text(file_path, file_type)
            elif file_type == "pdf":
                return self._process_pdf(file_path)
            elif file_type == "docx":
                return self._process_docx(file_path)
            elif file_type == "image":
                return self._process_image(file_path)
            elif file_type == "unknown":
                # Try to read as text anyway
                return self._process_text(file_path, "unknown")
            else:
                return ProcessingResult(
                    text="",
                    file_type=file_type,
                    success=False,
                    error=f"Unsupported file type: {file_type}",
                )
        except Exception as e:
            return ProcessingResult(
                text="",
                file_type=file_type,
                success=False,
                error=str(e),
            )

    def _process_text(self, file_path: Path, file_type: str) -> ProcessingResult:
        """Process a text/code file."""
        # Read raw bytes first
        raw_bytes = file_path.read_bytes()

        # Detect encoding
        detected = chardet.detect(raw_bytes)
        encoding = detected.get("encoding", "utf-8") or "utf-8"

        try:
            text = raw_bytes.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            # Fallback to utf-8 with error handling
            text = raw_bytes.decode("utf-8", errors="replace")

        return ProcessingResult(
            text=text,
            file_type=file_type,
            success=True,
        )

    def _process_pdf(self, file_path: Path) -> ProcessingResult:
        """Process a PDF file."""
        if not HAS_PDF:
            return ProcessingResult(
                text="",
                file_type="pdf",
                success=False,
                error="pypdf not installed. Install with: pip install pypdf",
            )

        reader = PdfReader(str(file_path))
        text_parts = []

        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

        return ProcessingResult(
            text="\n".join(text_parts),
            file_type="pdf",
            success=True,
        )

    def _process_docx(self, file_path: Path) -> ProcessingResult:
        """Process a DOCX file."""
        if not HAS_DOCX:
            return ProcessingResult(
                text="",
                file_type="docx",
                success=False,
                error="python-docx not installed. Install with: pip install python-docx",
            )

        doc = DocxDocument(str(file_path))
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text:
                text_parts.append(paragraph.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text_parts.append(cell.text)

        return ProcessingResult(
            text="\n".join(text_parts),
            file_type="docx",
            success=True,
        )

    def _process_image(self, file_path: Path) -> ProcessingResult:
        """Process an image file using OCR."""
        if not HAS_OCR:
            return ProcessingResult(
                text="",
                file_type="image",
                success=False,
                error="OCR not available. Install pytesseract and Pillow, and ensure Tesseract is installed.",
            )

        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)

            return ProcessingResult(
                text=text,
                file_type="image",
                success=True,
            )
        except pytesseract.TesseractNotFoundError:
            return ProcessingResult(
                text="",
                file_type="image",
                success=False,
                error="Tesseract not found. Please install Tesseract OCR.",
            )
